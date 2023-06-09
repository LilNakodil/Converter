import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE



# Предикат на вопрос, просто ищет слово "ответ в параграфах"
def isQuestion(text):
    word="ответа."
    return (word in text)

# Соединяет все раны параграфа в один, ибо библиотека может неадекватно выделять стили документа
def joiner(text):
    FullText=""
    for i in text.runs:
        FullText+=i.text
    return(FullText)

# на всякий случай переводит документ в один стиль
def styler(file):
    for p in file.paragraphs:
        if(isQuestion(p.text)!=True):
            p.style="Normal"
    
class questionn:
    # Класс вопроса, содержит массив правильных и неправильных ответов, и сам вопрос
    def __init__(self):
        self.question=""
        self.goodAnswer=[]
        self.badAnswer=[]
    # Чистка структуры
    def clear(self):
        self.question=""
        self.goodAnswer.clear
        self.badAnswer.clear
    # Перевод структуры в шаблонную строку
    def toString(self):
        string=""
        string+=self.question+"\n"
        string+="{ \n"
        if(len(self.goodAnswer)>1):  
            goodPercent=100//(len(self.goodAnswer))
            #badPercent=100//(len(self.badAnswer))
            for i in self.goodAnswer:
                string+="~%"+str(goodPercent)+"% "+i+"\n"
            for j in self.badAnswer:
                string+="~%-"+str(100)+"% "+j+"\n"
        else:
            string+="= "+ self.goodAnswer[0]+"\n"
            for j in self.badAnswer:
                string+="~ "+j+"\n"
        string+="} \n"
        return string
    
# Вывод в консоль строк массива с вопросами
def printStr(arr):
    for i in arr:
        print(i)

# Конвертация docx файла
def converter(file):
    IncorrectAnswers=[]
    correctAnswers=[]
    questions=[]
    adress=0
    # Идем по параграфам докса
    for p in file.paragraphs:
        # Стиль на всякий случай
        p.style="Normal"

        if(isQuestion(p.text)):
            #print(p.text+"\n{\n")
            adress+=1
            # Добавление вопроса и его id
            questions.append([p.text,adress])
            
        else:
            # Если мы сейчас не на вопросе, то идем по параграфу и смотрим структуру шрифта через runs
            for i in p.runs:
                # Если параграф не жирный и не обозначен курсивом, то это неправильный ответ (отталкиваясь от документа заказчика), пишем его в массив неправильных ответов с айдишником
                if(not(i.bold) and not(i.italic) and i.text!=" "):
                    IncorrectAnswers.append([joiner(p),adress])
                    #print("Incorrect " +joiner(p))
                    break
                # Если параграф  жирный и обозначен курсивом, то это правильный ответ, пишем его в массив правильных ответов с айдишником
                elif(i.bold and i.italic and i.text!=" "):
                    correctAnswers.append([joiner(p),adress])
                    #print("Correct " +joiner(p))
                    break
    # Заводим массив со строковыми вопросами
    quess=[]
    for adresses in range(0,len(questions)):
        que=questionn()
        que.question=questions[adresses][0]
        # Соотносим айдишник вопроса с правильными и неправильными ответами
        for j in range(0,len(correctAnswers)):
            if(questions[adresses][1]==correctAnswers[j][1]):
                que.goodAnswer.append(correctAnswers[j][0])
        for m in range(0,len(IncorrectAnswers)):
            if(questions[adresses][1]==IncorrectAnswers[m][1]):
                que.badAnswer.append(IncorrectAnswers[m][0])
        # Добавляем структуру в уже преобразованном виде
        quess.append(que.toString())
    return quess

# Перевод строк в txt    
def toTxt(testArray):
    output=open("converted.txt","w",encoding="utf-8")
    for i in testArray:
        output.write(i)
    output.close()

def mainn(): 
    doc=docx.Document("aboba.docx")     
    styler(doc)
    questions=converter(doc) 
    toTxt(questions)   

mainn()               


        
